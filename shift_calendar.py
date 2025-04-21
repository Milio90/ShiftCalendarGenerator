import docx
from icalendar import Calendar, Event
from datetime import datetime, timedelta
import re
import sys
import os
import tkinter as tk
from tkinter import filedialog

def browse_file():
    """Open a file browser dialog and return the selected file path."""
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    file_path = filedialog.askopenfilename(
        title="Select Shift Schedule Document",
        filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
    )
    return file_path

def read_docx_table(file_path):
    """Read the table content from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        # Assuming the first table contains our data
        if not doc.tables:
            print("No tables found in the document.")
            return []
        
        table = doc.tables[0]
        rows = []
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            # Skip empty rows
            if any(row_data):
                rows.append(row_data)
        
        return rows
    except Exception as e:
        print(f"Error reading document: {e}")
        return []

def parse_shifts(rows, month, year):
    """Parse the shift data from table rows."""
    shifts = []
    
    for row in rows:
        if len(row) < 4:  # Ensure row has enough columns
            continue
        
        try:
            # Extract day, month, day_of_week, and employees
            day = row[0].strip()
            day_of_week = row[2].strip()
            employees_cell = row[3].strip()
            
            # Skip header rows or rows without day number
            if not day.isdigit():
                continue
            
            day = int(day)
            
            # Parse employee names (may contain two employees, one with asterisk)
            employees = employees_cell.split('\n')
            employees = [e.strip() for e in employees if e.strip()]
            
            for employee in employees:
                is_on_call = "*" in employee
                employee_name = employee.replace("*", "").strip()
                
                # Create shift date
                shift_date = datetime(year, month, day)
                
                # A 24-hour shift typically starts in the morning and ends the next morning
                start_time = datetime(year, month, day, 8, 0, 0)  # 8:00 AM
                end_time = start_time + timedelta(hours=24)  # 24 hours later
                
                shift_type = "On-Call Shift" if is_on_call else "Regular Shift"
                
                shifts.append({
                    'employee': employee_name,
                    'date': shift_date,
                    'day_of_week': day_of_week,
                    'shift_type': shift_type,
                    'start_time': start_time,
                    'end_time': end_time
                })
        except Exception as e:
            print(f"Error parsing row {row}: {e}")
            continue
    
    return shifts

def create_calendar_for_employee(shifts, employee_name, output_file):
    """Create an iCalendar file with events for a specific employee."""
    employee_shifts = [s for s in shifts if s['employee'].lower() == employee_name.lower()]
    
    if not employee_shifts:
        print(f"No shifts found for employee: {employee_name}")
        return None
    
    cal = Calendar()
    cal.add('prodid', '-//Employee Shift Calendar//example.com//')
    cal.add('version', '2.0')
    
    for shift in employee_shifts:
        event = Event()
        
        # Set event properties
        summary = f"{shift['shift_type']} - {shift['day_of_week']}"
        event.add('summary', summary)
        event.add('dtstart', shift['start_time'])
        event.add('dtend', shift['end_time'])
        event.add('dtstamp', datetime.now())
        
        # Generate a unique ID for the event
        uid = f"{employee_name.replace(' ', '')}-{shift['date'].strftime('%Y%m%d')}@shifts.example.com"
        event.add('uid', uid)
        
        # Add description
        description = f"24-hour {shift['shift_type'].lower()} for {employee_name}"
        event.add('description', description)
        
        cal.add_component(event)
    
    # Write to file
    with open(output_file, 'wb') as f:
        f.write(cal.to_ical())
    
    return output_file

def save_calendar_file(employee_name):
    """Open a file dialog to choose where to save the calendar file."""
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    default_filename = f"{employee_name.replace(' ', '_')}_shifts.ics"
    file_path = filedialog.asksaveasfilename(
        title="Save Calendar File",
        defaultextension=".ics",
        filetypes=[("Calendar Files", "*.ics"), ("All Files", "*.*")],
        initialfile=default_filename
    )
    return file_path

def extract_month_year_from_filename(filename):
    """Attempt to extract month and year from the filename."""
    # Example: "ΕΦΗΜΕΡΙΕΣ ΜΑΡΤΙΟΣ 2025.docx"
    month_dict = {
        "ΙΑΝΟΥΑΡΙΟΣ": 1, "ΦΕΒΡΟΥΑΡΙΟΣ": 2, "ΜΑΡΤΙΟΣ": 3, "ΑΠΡΙΛΙΟΣ": 4,
        "ΜΑΙΟΣ": 5, "ΙΟΥΝΙΟΣ": 6, "ΙΟΥΛΙΟΣ": 7, "ΑΥΓΟΥΣΤΟΣ": 8,
        "ΣΕΠΤΕΜΒΡΙΟΣ": 9, "ΟΚΤΩΒΡΙΟΣ": 10, "ΝΟΕΜΒΡΙΟΣ": 11, "ΔΕΚΕΜΒΡΙΟΣ": 12
    }
    
    # Default to current month and year if extraction fails
    default_month = datetime.now().month
    default_year = datetime.now().year
    
    try:
        # Try to extract month name and year
        for month_name, month_num in month_dict.items():
            if month_name in filename:
                # Found month, now look for year
                year_match = re.search(r'20\d\d', filename)
                if year_match:
                    year = int(year_match.group())
                    return month_num, year
                return month_num, default_year
    except:
        pass
    
    return default_month, default_year

def main():
    print("Employee Shift Calendar Generator")
    print("=================================")
    
    # Get input file using file browser
    print("Please select the DOCX file containing shift schedules...")
    input_file = browse_file()
    
    if not input_file:
        print("No file selected. Exiting.")
        return
    
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        return
    
    print(f"Selected file: {input_file}")
    
    # Extract month and year from filename if possible
    month, year = extract_month_year_from_filename(os.path.basename(input_file))
    
    # Allow user to override detected month/year
    print(f"Detected: Month {month}, Year {year}")
    month_input = input(f"Enter month number (1-12) [default: {month}]: ").strip()
    if month_input and month_input.isdigit() and 1 <= int(month_input) <= 12:
        month = int(month_input)
    
    year_input = input(f"Enter year [default: {year}]: ").strip()
    if year_input and year_input.isdigit() and len(year_input) == 4:
        year = int(year_input)
    
    # Read and parse the document
    print(f"Reading file: {input_file}")
    rows = read_docx_table(input_file)
    
    if not rows:
        print("No data found in the document.")
        return
    
    print("Parsing shifts...")
    shifts = parse_shifts(rows, month, year)
    
    if not shifts:
        print("No shifts found in the document!")
        return
    
    # Get unique employee names
    all_employees = sorted(set(shift['employee'] for shift in shifts))
    print(f"Found {len(shifts)} shift assignments for {len(all_employees)} employees:")
    
    # Display employee list
    for i, emp in enumerate(all_employees, 1):
        print(f"{i}. {emp}")
    
    # Ask user which employee to generate calendar for
    while True:
        employee_choice = input("\nEnter employee name or number (or 'all' for all employees): ").strip()
        
        if employee_choice.lower() == 'all':
            # Generate calendars for all employees
            for employee in all_employees:
                output_file = save_calendar_file(employee)
                if output_file:
                    result = create_calendar_for_employee(shifts, employee, output_file)
                    if result:
                        print(f"Calendar for {employee} created successfully: {output_file}")
            break
        
        # Check if user entered a number
        elif employee_choice.isdigit() and 1 <= int(employee_choice) <= len(all_employees):
            employee_name = all_employees[int(employee_choice) - 1]
        else:
            # Assume user entered a name
            employee_name = employee_choice
            # Check if name exists in our list
            if employee_name.lower() not in [emp.lower() for emp in all_employees]:
                closest_match = None
                for emp in all_employees:
                    if employee_name.lower() in emp.lower():
                        closest_match = emp
                        break
                
                if closest_match:
                    confirm = input(f"Did you mean '{closest_match}'? (y/n): ").lower()
                    if confirm == 'y':
                        employee_name = closest_match
                    else:
                        print("Please try again.")
                        continue
                else:
                    print(f"Employee '{employee_name}' not found. Please try again.")
                    continue
        
        # Get output file location using file browser
        print(f"Select where to save the calendar file for {employee_name}...")
        output_file = save_calendar_file(employee_name)
        
        if not output_file:
            print("Calendar save operation cancelled.")
            retry = input("Would you like to select a different employee? (y/n): ").lower()
            if retry == 'y':
                continue
            else:
                break
        
        # Create calendar
        result = create_calendar_for_employee(shifts, employee_name, output_file)
        if result:
            print(f"Calendar created successfully: {output_file}")
            
        # Ask if user wants to create another calendar
        another = input("Would you like to create a calendar for another employee? (y/n): ").lower()
        if another != 'y':
            break

if __name__ == "__main__":
    main()
