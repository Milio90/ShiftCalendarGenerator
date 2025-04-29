import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from datetime import datetime, timedelta, date
import os
import re
import threading
import docx
from icalendar import Calendar, Event
import platform
import subprocess
import tempfile
import shutil

class ShiftCalendarApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Employee Shift Calendar Generator")
        self.root.geometry("800x600")
        self.root.minsize(800, 600)
        
        # Variables
        self.input_file = tk.StringVar()
        self.cath_lab_file = tk.StringVar()
        self.ep_file = tk.StringVar()
        self.month = tk.IntVar(value=datetime.now().month)
        self.year = tk.IntVar(value=datetime.now().year)
        self.include_cath_lab = tk.BooleanVar(value=False)
        self.include_ep = tk.BooleanVar(value=False)
        
        # Data storage
        self.all_shifts = []
        self.cath_lab_shifts = []
        self.ep_shifts = []
        self.all_employees = []
        
        # Create UI
        self.create_widgets()
        
        # Configure grid weights
        self.root.grid_columnconfigure(0, weight=1)
        for i in range(6):
            self.root.grid_rowconfigure(i, weight=0)
        self.root.grid_rowconfigure(6, weight=1)  # Log area takes remaining space
        
    def create_widgets(self):
        # File selection frame
        file_frame = ttk.LabelFrame(self.root, text="File Selection")
        file_frame.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        file_frame.grid_columnconfigure(1, weight=1)
        
        # Main shifts file
        ttk.Label(file_frame, text="Main Shift Schedule:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        ttk.Entry(file_frame, textvariable=self.input_file, width=50).grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ttk.Button(file_frame, text="Browse", command=self.browse_main_file).grid(row=0, column=2, padx=5, pady=5)
        
        # Month and Year frame
        date_frame = ttk.LabelFrame(self.root, text="Schedule Date")
        date_frame.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        
        ttk.Label(date_frame, text="Month:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        month_combo = ttk.Combobox(date_frame, textvariable=self.month, width=15)
        month_combo['values'] = list(range(1, 13))
        month_combo.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        ttk.Label(date_frame, text="Year:").grid(row=0, column=2, padx=5, pady=5, sticky="w")
        year_combo = ttk.Combobox(date_frame, textvariable=self.year, width=10)
        current_year = datetime.now().year
        year_combo['values'] = list(range(current_year-1, current_year+3))
        year_combo.grid(row=0, column=3, padx=5, pady=5, sticky="w")
        
        # Optional schedules frame
        opt_frame = ttk.LabelFrame(self.root, text="Optional Schedules")
        opt_frame.grid(row=2, column=0, padx=10, pady=5, sticky="ew")
        opt_frame.grid_columnconfigure(1, weight=1)
        
        # Cath Lab file
        ttk.Checkbutton(opt_frame, text="Include Cath Lab On-Call", variable=self.include_cath_lab, 
                      command=self.toggle_cath_lab).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.cath_entry = ttk.Entry(opt_frame, textvariable=self.cath_lab_file, width=50, state="disabled")
        self.cath_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        self.cath_button = ttk.Button(opt_frame, text="Browse", command=self.browse_cath_file, state="disabled")
        self.cath_button.grid(row=0, column=2, padx=5, pady=5)
        
        # EP file
        ttk.Checkbutton(opt_frame, text="Include Electrophysiology On-Call", variable=self.include_ep,
                      command=self.toggle_ep).grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.ep_entry = ttk.Entry(opt_frame, textvariable=self.ep_file, width=50, state="disabled")
        self.ep_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        self.ep_button = ttk.Button(opt_frame, text="Browse", command=self.browse_ep_file, state="disabled")
        self.ep_button.grid(row=1, column=2, padx=5, pady=5)
        
        # Process button
        ttk.Button(self.root, text="Process Files", command=self.process_files).grid(row=3, column=0, padx=10, pady=10)
        
        # Employee selection frame (initially hidden)
        self.employee_frame = ttk.LabelFrame(self.root, text="Employee Selection")
        self.employee_frame.grid(row=4, column=0, padx=10, pady=5, sticky="ew")
        self.employee_frame.grid_remove()  # Hide initially
        
        # Employee listbox
        self.employee_listbox = tk.Listbox(self.employee_frame, height=5, selectmode=tk.EXTENDED)
        self.employee_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Buttons for employee selection
        button_frame = ttk.Frame(self.employee_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(button_frame, text="Generate Selected", command=self.generate_selected).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Generate All", command=self.generate_all).pack(side=tk.LEFT, padx=5)
        
        # Status/log area
        log_frame = ttk.LabelFrame(self.root, text="Status Log")
        log_frame.grid(row=6, column=0, padx=10, pady=5, sticky="nsew")
        log_frame.grid_columnconfigure(0, weight=1)
        log_frame.grid_rowconfigure(0, weight=1)
        
        self.log_area = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, height=10)
        self.log_area.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
        self.log_area.config(state=tk.DISABLED)
        
    def toggle_cath_lab(self):
        if self.include_cath_lab.get():
            self.cath_entry.config(state="normal")
            self.cath_button.config(state="normal")
        else:
            self.cath_entry.config(state="disabled")
            self.cath_button.config(state="disabled")
            
    def toggle_ep(self):
        if self.include_ep.get():
            self.ep_entry.config(state="normal")
            self.ep_button.config(state="normal")
        else:
            self.ep_entry.config(state="disabled")
            self.ep_button.config(state="disabled")
            
    def browse_main_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Main Shift Schedule Document",
            filetypes=[("Word Documents", "*.docx *.doc"), ("All Files", "*.*")]
        )
        if file_path:
            self.input_file.set(file_path)
            self.log(f"Selected main file: {file_path}")
            
            # Try to extract month and year from filename
            month, year = self.extract_month_year_from_filename(os.path.basename(file_path))
            self.month.set(month)
            self.year.set(year)
            self.log(f"Detected: Month {month}, Year {year}")
            
    def browse_cath_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Cath Lab On-Call Schedule",
            filetypes=[("Word Documents", "*.docx *.doc"), ("All Files", "*.*")]
        )
        if file_path:
            self.cath_lab_file.set(file_path)
            self.log(f"Selected Cath Lab file: {file_path}")
            
    def browse_ep_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Electrophysiology On-Call Schedule",
            filetypes=[("Word Documents", "*.docx *.doc"), ("All Files", "*.*")]
        )
        if file_path:
            self.ep_file.set(file_path)
            self.log(f"Selected Electrophysiology file: {file_path}")
            
    def log(self, message):
        """Add message to the log area"""
        self.log_area.config(state=tk.NORMAL)
        self.log_area.insert(tk.END, message + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state=tk.DISABLED)
        print(message)  # Also print to console for debugging
        
    def process_files(self):
        """Process the selected files and populate employee list"""
        input_file = self.input_file.get()
        
        if not input_file:
            messagebox.showerror("Error", "Please select a main shift schedule file.")
            return
            
        if not os.path.exists(input_file):
            messagebox.showerror("Error", f"File not found: {input_file}")
            return
            
        # Start processing in a separate thread to prevent UI freezing
        threading.Thread(target=self._process_files_thread, daemon=True).start()
        
    def _process_files_thread(self):
        """Thread function to process files"""
        # Clear previous data
        self.all_shifts = []
        self.cath_lab_shifts = []
        self.ep_shifts = []
        self.all_employees = []
        
        # Process main file
        input_file = self.input_file.get()
        self.log(f"Processing main file: {input_file}")
        
        month = self.month.get()
        year = self.year.get()
        
        # Process Cath Lab file if selected
        if self.include_cath_lab.get():
            cath_lab_file = self.cath_lab_file.get()
            if cath_lab_file and os.path.exists(cath_lab_file):
                self.log(f"Processing Cath Lab file: {cath_lab_file}")
                cath_lab_tables = self.read_docx_tables(cath_lab_file)
                if cath_lab_tables:
                    for table in cath_lab_tables:
                        cath_shifts = self.parse_specialty_on_call_table(table)
                        for shift in cath_shifts:
                            shift['shift_type'] = "Cath Lab On-Call"
                        self.cath_lab_shifts.extend(cath_shifts)
                    self.log(f"Found {len(self.cath_lab_shifts)} Cath Lab on-call shifts")
                else:
                    self.log("No tables found in the Cath Lab schedule document.")
            else:
                self.log("Cath Lab file not selected or not found.")
                
        # Process EP file if selected
        if self.include_ep.get():
            ep_file = self.ep_file.get()
            if ep_file and os.path.exists(ep_file):
                self.log(f"Processing Electrophysiology file: {ep_file}")
                ep_tables = self.read_docx_tables(ep_file)
                if ep_tables:
                    for table in ep_tables:
                        electro_shifts = self.parse_specialty_on_call_table(table)
                        for shift in electro_shifts:
                            shift['shift_type'] = "Electrophysiology On-Call"
                        self.ep_shifts.extend(electro_shifts)
                    self.log(f"Found {len(self.ep_shifts)} Electrophysiology on-call shifts")
                else:
                    self.log("No tables found in the Electrophysiology schedule document.")
            else:
                self.log("Electrophysiology file not selected or not found.")
                
        # Process main file tables
        tables = self.read_docx_tables(input_file)
        
        if not tables:
            self.log("No tables found in the main document.")
            return
            
        # Parse shifts from tables
        if len(tables) >= 1:
            self.log("Parsing first table (Regular/On-Call shifts)...")
            first_table_shifts = self.parse_first_table(tables[0], month, year)
            self.all_shifts.extend(first_table_shifts)
            self.log(f"Found {len(first_table_shifts)} shifts in first table")
            
        if len(tables) >= 2:
            self.log("Parsing second table (Μεγάλη/Μικρή/ΤΕΠ shifts)...")
            second_table_shifts = self.parse_second_table(tables[1], month, year)
            self.all_shifts.extend(second_table_shifts)
            self.log(f"Found {len(second_table_shifts)} shifts in second table")
            
        if not self.all_shifts:
            self.log("No shifts found in any table!")
            return
            
        # Get unique employee names
        self.all_employees = sorted(set(shift['employee'] for shift in self.all_shifts))
        self.log(f"Found {len(self.all_shifts)} total shift assignments for {len(self.all_employees)} employees")
        
        # Update UI with employee list (in main thread)
        self.root.after(0, self.update_employee_list)
        
    def update_employee_list(self):
        """Update the employee listbox with found employees"""
        # Show the employee frame
        self.employee_frame.grid()
        
        # Clear and populate listbox
        self.employee_listbox.delete(0, tk.END)
        for emp in self.all_employees:
            self.employee_listbox.insert(tk.END, emp)
            
        self.log("Please select employee(s) to generate calendar for")
            
    def generate_selected(self):
        """Generate calendars for selected employees"""
        selected_indices = self.employee_listbox.curselection()
        if not selected_indices:
            messagebox.showinfo("Information", "Please select at least one employee.")
            return
            
        selected_employees = [self.employee_listbox.get(i) for i in selected_indices]
        self.generate_calendars(selected_employees)
        
    def generate_all(self):
        """Generate calendars for all employees"""
        self.generate_calendars(self.all_employees)
        
    def generate_calendars(self, employees):
        """Generate calendar files for the specified employees"""
        if not employees:
            return
            
        # Ask for output directory
        output_dir = filedialog.askdirectory(title="Select Output Directory for Calendar Files")
        if not output_dir:
            self.log("Calendar generation cancelled - no output directory selected")
            return
            
        self.log(f"Generating calendars for {len(employees)} employees...")
        
        # Start generation in a separate thread
        threading.Thread(target=self._generate_calendars_thread, 
                         args=(employees, output_dir), 
                         daemon=True).start()
            
    def _generate_calendars_thread(self, employees, output_dir):
        """Thread function to generate calendar files"""
        success_count = 0
        
        for employee in employees:
            output_file = os.path.join(output_dir, f"{employee.replace(' ', '_')}_shifts.ics")
            
            # Create calendar
            result = self.create_calendar_for_employee(
                self.all_shifts, 
                employee, 
                output_file, 
                self.cath_lab_shifts, 
                self.ep_shifts
            )
            
            if result:
                success_count += 1
                self.log(f"Calendar for {employee} created successfully")
            else:
                self.log(f"Failed to create calendar for {employee}")
                
        self.log(f"Completed! Generated {success_count} of {len(employees)} calendars in {output_dir}")
        
        # Show completion message
        self.root.after(0, lambda: messagebox.showinfo(
            "Complete", 
            f"Generated {success_count} of {len(employees)} calendars."
        ))

    # Core functionality methods (adapted from original code)
    def extract_month_year_from_filename(self, filename):
        """Attempt to extract month and year from the filename."""
        # Example: "ΕΦΗΜΕΡΙΕΣ ΜΑΡΤΙΟΣ 2025.docx"
        month_dict = {
            "ΙΑΝΟΥΑΡΙΟΣ": 1, "ΦΕΒΡΟΥΑΡΙΟΣ": 2, "ΜΑΡΤΙΟΣ": 3, "ΑΠΡΙΛΙΟΣ": 4,
            "ΜΑΙΟΣ": 5, "ΙΟΥΝΙΟΣ": 6, "ΙΟΥΛΙΟΣ": 7, "ΑΥΓΟΥΣΤΟΣ": 8,
            "ΣΕΠΤΕΜΒΡΙΟΣ": 9, "ΟΚΤΩΒΡΙΟΣ": 10, "ΝΟΕΜΒΡΙΟΣ": 11, "ΔΕΚΕΜΒΡΙΟΣ": 12
        }
        
        # Default to current month and year if extraction fails
        default_month = datetime.now().month
        default_year = datetime.now().year
        
        try:
            # Try to extract month name and year
            for month_name, month_num in month_dict.items():
                if month_name in filename:
                    # Found month, now look for year
                    year_match = re.search(r'20\d\d', filename)
                    if year_match:
                        year = int(year_match.group())
                        return month_num, year
                    return month_num, default_year
                
        except:
            pass
        
        return default_month, default_year

    def convert_doc_to_docx(self, doc_path):
        """Convert a .doc file to .docx format using available tools."""
        file_name, file_ext = os.path.splitext(doc_path)
        
        # If already a docx file, return the original path
        if file_ext.lower() == '.docx':
            return doc_path
        
        # Create a temporary output file
        temp_dir = tempfile.gettempdir()
        base_name = os.path.basename(file_name)
        output_path = os.path.join(temp_dir, f"{base_name}_converted.docx")
        
        conversion_successful = False
        error_message = ""
        
        # Try LibreOffice first (cross-platform)
        try:
            # Determine LibreOffice executable based on platform
            libreoffice_cmd = None
            if platform.system() == "Windows":
                # Check common install locations
                possible_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                for path in possible_paths:
                    if os.path.exists(path):
                        libreoffice_cmd = path
                        break
            elif platform.system() == "Darwin":  # macOS
                libreoffice_cmd = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            else:  # Linux and others
                libreoffice_cmd = "libreoffice"
            
            if libreoffice_cmd:
                process = subprocess.run([
                    libreoffice_cmd,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", temp_dir,
                    doc_path
                ], capture_output=True, text=True, timeout=30)
                
                # LibreOffice sometimes creates with original filename in the output dir
                expected_file = os.path.join(temp_dir, f"{base_name}.docx")
                if os.path.exists(expected_file):
                    # Rename to our expected output path
                    shutil.move(expected_file, output_path)
                    conversion_successful = True
                else:
                    error_message = f"LibreOffice conversion output file not found: {expected_file}"
        except Exception as e:
            error_message = f"LibreOffice conversion failed: {str(e)}"
        
        # If LibreOffice failed, try Microsoft Word automation (Windows only)
        if not conversion_successful and platform.system() == "Windows":
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(output_path, FileFormat=16)  # 16 = docx format
                doc.Close()
                word.Quit()
                
                if os.path.exists(output_path):
                    conversion_successful = True
                else:
                    error_message += "\nMicrosoft Word conversion output file not found."
            except Exception as e:
                error_message += f"\nMicrosoft Word conversion failed: {str(e)}"
        
        if conversion_successful:
            self.log(f"Successfully converted {doc_path} to {output_path}")
            return output_path
        else:
            self.log(f"Failed to convert .doc to .docx: {error_message}")
            raise Exception(f"Could not convert {doc_path} to .docx format. Please convert it manually and try again.")

    def read_docx_tables(self, file_path):
        """Read all tables content from a DOCX file."""
        try:
            # Check if file is .doc and convert if needed
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext == '.doc':
                self.log("Converting .doc file to .docx format...")
                file_path = self.convert_doc_to_docx(file_path)
            
            doc = docx.Document(file_path)
            if not doc.tables:
                self.log("No tables found in the document.")
                return []
            
            tables_data = []
            
            for table_index, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    # Skip empty rows
                    if any(row_data):
                        rows.append(row_data)
                
                tables_data.append(rows)
                self.log(f"Table {table_index+1}: Found {len(rows)} rows with data")
            
            return tables_data
        except Exception as e:
            self.log(f"Error reading document: {e}")
            if "Could not convert" in str(e):
                # This is our custom error from conversion function
                self.log(str(e))
            messagebox.showerror("Error", f"Could not process the document: {e}")
            return []

    def parse_first_table(self, rows, month, year):
        """Parse the first table format (Regular and On-Call shifts) with month rollover detection."""
        shifts = []
        current_month = month
        current_year = year
        last_day = 0  # Track the last day number we've seen
        
        # Dictionary of Greek month names to month numbers
        greek_months = {
            "ΙΑΝΟΥΑΡΙΟΥ": 1, "ΦΕΒΡΟΥΑΡΙΟΥ": 2, "ΜΑΡΤΙΟΥ": 3, "ΑΠΡΙΛΙΟΥ": 4,
            "ΜΑΙΟΥ": 5, "ΙΟΥΝΙΟΥ": 6, "ΙΟΥΛΙΟΥ": 7, "ΑΥΓΟΥΣΤΟΥ": 8,
            "ΣΕΠΤΕΜΒΡΙΟΥ": 9, "ΟΚΤΩΒΡΙΟΥ": 10, "ΝΟΕΜΒΡΙΟΥ": 11, "ΔΕΚΕΜΒΡΙΟΥ": 12,
            "ΙΑΝΟΥΑΡΙΟΣ": 1, "ΦΕΒΡΟΥΑΡΙΟΣ": 2, "ΜΑΡΤΙΟΣ": 3, "ΑΠΡΙΛΙΟΣ": 4,
            "ΜΑΙΟΣ": 5, "ΙΟΥΝΙΟΣ": 6, "ΙΟΥΛΙΟΣ": 7, "ΑΥΓΟΥΣΤΟΣ": 8,
            "ΣΕΠΤΕΜΒΡΙΟΣ": 9, "ΟΚΤΩΒΡΙΟΣ": 10, "ΝΟΕΜΒΡΙΟΣ": 11, "ΔΕΚΕΜΒΡΙΟΣ": 12
        }
        
        for row in rows:
            if len(row) < 4:  # Ensure row has enough columns
                continue
            
            try:
                # Extract day, month_text, day_of_week, and employees
                day = row[0].strip()
                month_text = row[1].strip() if len(row) > 1 else ""
                day_of_week = row[2].strip()
                employees_cell = row[3].strip()
                
                # Skip header rows or rows without day number
                if not day or not day[0].isdigit():
                    continue
                
                # Handle special formatting like "*01**" for May 1st
                day = day.strip("*").strip()
                if not day.isdigit():
                    continue
                    
                day = int(day)
                
                # Check for explicit month name in the month_text field
                found_month = None
                for greek_month, month_num in greek_months.items():
                    if greek_month in month_text:
                        found_month = month_num
                        break
                
                if found_month is not None:
                    # Use explicitly mentioned month
                    current_month = found_month
                    # If the new month is less than the original month, we've moved to next year
                    if current_month < month and month > 10 and current_month < 3:
                        current_year += 1
                    self.log(f"Explicit month found: now processing {current_month}/{current_year}")
                elif day < last_day and last_day > 20 and day < 10:
                    # Move to next month based on day number patterns
                    current_month += 1
                    if current_month > 12:
                        current_month = 1
                        current_year += 1
                    self.log(f"Month rollover detected: now processing {current_month}/{current_year}")
                
                last_day = day
                
                # Parse employee names (may contain two employees, one with asterisk)
                employees = employees_cell.split('\n')
                employees = [e.strip() for e in employees if e.strip()]
                
                for employee in employees:
                    is_on_call = "*" in employee
                    employee_name = employee.replace("*", "").strip()
                    
                    # Create shift date using current_month and current_year
                    shift_date = date(current_year, current_month, day)
                    
                    shift_type = "On-Call Shift" if is_on_call else "Regular Shift"
                    
                    shifts.append({
                        'employee': employee_name,
                        'date': shift_date,
                        'day_of_week': day_of_week,
                        'shift_type': shift_type
                    })
            except Exception as e:
                self.log(f"Error parsing row in first table {row}: {e}")
                continue
        
        return shifts

    def parse_second_table(self, rows, month, year):
        """Parse the second table format (Μεγάλη, Μικρή, ΤΕΠ shifts) with month rollover detection."""
        shifts = []
        current_month = month
        current_year = year
        last_day = 0  # Track the last day number we've seen
        
        # Dictionary of Greek month names to month numbers
        greek_months = {
            "ΙΑΝΟΥΑΡΙΟΥ": 1, "ΦΕΒΡΟΥΑΡΙΟΥ": 2, "ΜΑΡΤΙΟΥ": 3, "ΑΠΡΙΛΙΟΥ": 4,
            "ΜΑΙΟΥ": 5, "ΙΟΥΝΙΟΥ": 6, "ΙΟΥΛΙΟΥ": 7, "ΑΥΓΟΥΣΤΟΥ": 8,
            "ΣΕΠΤΕΜΒΡΙΟΥ": 9, "ΟΚΤΩΒΡΙΟΥ": 10, "ΝΟΕΜΒΡΙΟΥ": 11, "ΔΕΚΕΜΒΡΙΟΥ": 12,
            "ΙΑΝΟΥΑΡΙΟΣ": 1, "ΦΕΒΡΟΥΑΡΙΟΣ": 2, "ΜΑΡΤΙΟΣ": 3, "ΑΠΡΙΛΙΟΣ": 4,
            "ΜΑΙΟΣ": 5, "ΙΟΥΝΙΟΣ": 6, "ΙΟΥΛΙΟΣ": 7, "ΑΥΓΟΥΣΤΟΣ": 8,
            "ΣΕΠΤΕΜΒΡΙΟΣ": 9, "ΟΚΤΩΒΡΙΟΣ": 10, "ΝΟΕΜΒΡΙΟΣ": 11, "ΔΕΚΕΜΒΡΙΟΣ": 12
        }
        
        for row in rows:
            if len(row) < 6:  # Ensure row has enough columns for second table format
                continue
            
            try:
                # Extract day, month_text, day_of_week, and employees from different shifts
                day = row[0].strip()
                month_text = row[1].strip() if len(row) > 1 else ""
                day_of_week = row[2].strip()
                megali_shift = row[3].strip()
                mikri_shift = row[4].strip()
                tep_shift = row[5].strip()
                
                # Skip header rows or rows without day number
                if not day or not day[0].isdigit():
                    continue
                
                # Handle special formatting like "*01**" for May 1st
                day = day.strip("*").strip()
                if not day.isdigit():
                    continue
                    
                day = int(day)
                
                # Check for explicit month name in the month_text field
                found_month = None
                for greek_month, month_num in greek_months.items():
                    if greek_month in month_text:
                        found_month = month_num
                        break
                
                if found_month is not None:
                    # Use explicitly mentioned month
                    current_month = found_month
                    # If the new month is less than the original month, we've moved to next year
                    if current_month < month and month > 10 and current_month < 3:
                        current_year += 1
                    self.log(f"Explicit month found: now processing {current_month}/{current_year}")
                elif day < last_day and last_day > 20 and day < 10:
                    # Move to next month based on day number patterns
                    current_month += 1
                    if current_month > 12:
                        current_month = 1
                        current_year += 1
                    self.log(f"Month rollover detected: now processing {current_month}/{current_year}")
                
                last_day = day
                
                # Use current_month and current_year for the shift date
                shift_date = date(current_year, current_month, day)
                
                # Process Μεγάλη shift (24h)
                if megali_shift:
                    employee_name = megali_shift.replace(">", "").strip()
                    if employee_name:
                        shifts.append({
                            'employee': employee_name,
                            'date': shift_date,
                            'day_of_week': day_of_week,
                            'shift_type': "Μεγάλη Shift (24h)"
                        })
                
                # Process Μικρή shift (24h)
                if mikri_shift:
                    employee_name = mikri_shift.replace(">", "").strip()
                    if employee_name:
                        shifts.append({
                            'employee': employee_name,
                            'date': shift_date,
                            'day_of_week': day_of_week,
                            'shift_type': "Μικρή Shift (24h)"
                        })
                
                # Process ΤΕΠ shift (12h)
                if tep_shift:
                    employee_name = tep_shift.replace(">", "").strip()
                    if employee_name:
                        shifts.append({
                            'employee': employee_name,
                            'date': shift_date,
                            'day_of_week': day_of_week,
                            'shift_type': "TEP Shift (12h)"
                        })
                    
            except Exception as e:
                self.log(f"Error parsing row in second table {row}: {e}")
                continue
        
        return shifts

    def parse_specialty_on_call_table(self, rows):
        """Parse the specialty on-call table format with date (DD-MM-YYYY or DD/MM/YYYY) in first column."""
        shifts = []
        
        for row in rows:
            if len(row) < 3:  # Ensure row has enough columns
                continue
            
            try:
                # Extract date, day_of_week, and employee
                date_str = row[0].strip()
                day_of_week = row[1].strip()
                employee_name = row[2].strip()
                
                # Skip header rows or rows without proper date format
                # Updated regex to match both DD-MM-YYYY and DD/MM/YYYY formats
                if not re.match(r"\d{1,2}[-/]\d{1,2}[-/]\d{4}", date_str):
                    continue
                
                # Parse date (supports both DD-MM-YYYY and DD/MM/YYYY)
                if '-' in date_str:
                    day, month, year = map(int, date_str.split('-'))
                elif '/' in date_str:
                    day, month, year = map(int, date_str.split('/'))
                else:
                    continue  # Skip if date format doesn't match either pattern
                    
                shift_date = date(year, month, day)
                
                if employee_name:
                    shifts.append({
                        'employee': employee_name,
                        'date': shift_date,
                        'day_of_week': day_of_week,
                        'shift_type': "On-Call Specialty",  # Will be updated when adding to all_shifts
                    })
                    
            except Exception as e:
                self.log(f"Error parsing row in specialty on-call table {row}: {e}")
                continue
        
        return shifts

    def create_calendar_for_employee(self, shifts, employee_name, output_file, cath_lab_shifts=None, ep_shifts=None):
        """Create an iCalendar file with all-day events for a specific employee."""
        # Filter shifts for this specific employee
        employee_shifts = [s for s in shifts if s['employee'].lower() == employee_name.lower()]
        
        # Also check if the employee has any cath lab or EP shifts
        employee_cath_lab_shifts = []
        employee_ep_shifts = []
        
        if cath_lab_shifts:
            employee_cath_lab_shifts = [s for s in cath_lab_shifts if s['employee'].lower() == employee_name.lower()]
            
        if ep_shifts:
            employee_ep_shifts = [s for s in ep_shifts if s['employee'].lower() == employee_name.lower()]
        
        if not employee_shifts and not employee_cath_lab_shifts and not employee_ep_shifts:
            self.log(f"No shifts found for employee: {employee_name}")
            return None
        
        cal = Calendar()
        cal.add('prodid', '-//Employee Shift Calendar//example.com//')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        
        # Group shifts by date to combine multiple shifts on the same day
        shifts_by_date = {}
        
        # Add regular shifts to the grouping
        for shift in employee_shifts:
            date_key = shift['date'].isoformat()
            if date_key not in shifts_by_date:
                shifts_by_date[date_key] = []
            shifts_by_date[date_key].append(shift)
        
        # Add cath lab shifts if they don't overlap with existing dates
        for shift in employee_cath_lab_shifts:
            date_key = shift['date'].isoformat()
            if date_key not in shifts_by_date:
                shifts_by_date[date_key] = []
            shifts_by_date[date_key].append(shift)
        
        # Add EP shifts if they don't overlap with existing dates
        for shift in employee_ep_shifts:
            date_key = shift['date'].isoformat()
            if date_key not in shifts_by_date:
                shifts_by_date[date_key] = []
            shifts_by_date[date_key].append(shift)
        
        # Create events for each date, combining shift information
        for date_key, date_shifts in shifts_by_date.items():
            event = Event()
            
            # Combine all shift types for the summary
            shift_types = [s['shift_type'] for s in date_shifts]
            day_of_week = date_shifts[0]['day_of_week']  # They all have the same date
            shift_date = date_shifts[0]['date']
            
            # Format the summary to show all shift types
            summary = f"{', '.join(shift_types)} - {day_of_week}"
            event.add('summary', summary)
            
            # All-day events need a DATE value type
            event.add('dtstart', shift_date)
            
            # For all-day events, the end date should be the next day
            # The end date is non-inclusive in the iCalendar spec
            end_date = shift_date + timedelta(days=1)
            event.add('dtend', end_date)
            
            event.add('dtstamp', datetime.now())
            
            # Generate a unique ID for the event
            uid = f"{employee_name.replace(' ', '')}-{shift_date.strftime('%Y%m%d')}@shifts.example.com"
            event.add('uid', uid)
            
            # Add description with details about all employees working that day
            description_parts = [f"Your shifts: {', '.join(shift_types)}"]
            
            # Find all employees working on this date
            coworkers_info = []
            for s in shifts:
                # If it's the same date but not the current employee
                if s['date'] == shift_date and s['employee'].lower() != employee_name.lower():
                    coworkers_info.append(f"{s['employee']}: {s['shift_type']}")
            
            # Add coworkers section if any exist
            if coworkers_info:
                description_parts.append("\nCoworkers on this day:")
                for info in sorted(coworkers_info):
                    description_parts.append(f"- {info}")
            else:
                description_parts.append("\nNo other employees scheduled on this day.")
            
            # Add Cath Lab on-call information if available
            if cath_lab_shifts:
                cath_lab_employee = None
                for shift in cath_lab_shifts:
                    if shift['date'] == shift_date and shift['employee'].lower() != employee_name.lower():
                        cath_lab_employee = shift['employee']
                        break
                
                if cath_lab_employee:
                    description_parts.append(f"\nCath Lab On-Call: {cath_lab_employee}")
            
            # Add Electrophysiology on-call information if available
            if ep_shifts:
                ep_employee = None
                for shift in ep_shifts:
                    if shift['date'] == shift_date and shift['employee'].lower() != employee_name.lower():
                        ep_employee = shift['employee']
                        break
                
                if ep_employee:
                    description_parts.append(f"\nElectrophysiology On-Call: {ep_employee}")
            
            event.add('description', "\n".join(description_parts))
            
            cal.add_component(event)
        
        # Write to file
        try:
            with open(output_file, 'wb') as f:
                f.write(cal.to_ical())
            return output_file
        except Exception as e:
            self.log(f"Error saving calendar file: {e}")
            return None


def main():
    # Create root window
    root = tk.Tk()
    root.wm_class("ShiftCalendarGenerator")
    # Set theme (if ttk styles available)
    try:
        style = ttk.Style()
        if 'clam' in style.theme_names():
            style.theme_use('clam')
    except:
        pass
    
    app = ShiftCalendarApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
